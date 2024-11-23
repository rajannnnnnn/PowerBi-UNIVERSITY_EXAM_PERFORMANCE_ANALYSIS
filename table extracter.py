from docx import Document
import pandas as pd

def print_dictionary(the_big_dictionary):
    for key in the_big_dictionary.keys():
        print(key,end=", ")
    print()
    for i in range(len(the_big_dictionary["Register Number"])):
        print(the_big_dictionary["Register Number"][i],the_big_dictionary["Student Name"][i],
              the_big_dictionary["Subject Code"][i],the_big_dictionary["Grade"][i],the_big_dictionary["Exam Date"][i])
def convert_to_dictionary(file_location,model,exam_date = ""):
    doc = Document(file_location)
    if exam_date == "":    
        n=0
        catch = -1
        exam_date = list()
        remove_data = ["Examination","Provisional Results of",".",","]
        for paragraph in doc.paragraphs:
            if (paragraph.text == "ANNA UNIVERSITY :: CHENNAI - 600025."):
                catch =  3
                n+=1
            catch = catch - 1
            if catch == 0:
                text = paragraph.text
                for x in remove_data:
                    text = text.replace(x,"")
                exam_date.append(text[1:].replace(" / ","/"))
        if not len(exam_date) is len(doc.tables):
            raise Exception("Table and Title Matching Error")
    else:
        exam_date = [exam_date for i in range(len(doc.tables))]
    the_big_dictionary = {"Register Number":[],"Student Name":[],"Subject Code":[],"Grade":[],"Exam Date":[]}
    if model == "full":
        for table_index, table in enumerate(doc.tables):
            if table.rows[0].cells[0].text == "": 
                sub_codes = [cell.text for cell in table.rows[0].cells[3:]]
                grades = [row.cells[3:] for row in table.rows[2:]]
                for i,sub_code in enumerate(sub_codes):
                    grade_list = []
                    for j,grade in enumerate(grades):
                        grade_list.append(grade[i].text)
                    the_big_dictionary["Subject Code"].extend([sub_code for row in table.rows[2:]])
                    the_big_dictionary["Subject Code"].append(" ")
                    the_big_dictionary["Register Number"].extend([row.cells[0].text for row in table.rows[2:]])
                    the_big_dictionary["Register Number"].append(" ")
                    the_big_dictionary["Student Name"].extend([row.cells[1].text.replace("\n"," ") for row in table.rows[2:]])
                    the_big_dictionary["Student Name"].append(" ")
                    the_big_dictionary["Grade"].extend(grade_list)
                    the_big_dictionary["Grade"].append(" ")
                    the_big_dictionary["Exam Date"].extend([exam_date[table_index] for row in table.rows[2:]])
                    the_big_dictionary["Exam Date"].append(" ")
    if model=="result release":
        for table_index, table in enumerate(doc.tables):
            if table.rows[0].cells[0].text == "": 
                sub_codes = [cell.text for cell in table.rows[0].cells[2:]]
                grades = [row.cells[2:] for row in table.rows[2:]]
                for i,sub_code in enumerate(sub_codes):
                    grade_list = []
                    for j,grade in enumerate(grades):
                        grade_list.append(grade[i].text)
                    the_big_dictionary["Subject Code"].extend([sub_code for row in table.rows[2:]])
                    the_big_dictionary["Subject Code"].append(" ")
                    the_big_dictionary["Register Number"].extend([row.cells[0].text for row in table.rows[2:]])
                    the_big_dictionary["Register Number"].append(" ")
                    the_big_dictionary["Student Name"].extend([row.cells[1].text.replace("\n"," ") for row in table.rows[2:]])
                    the_big_dictionary["Student Name"].append(" ")
                    the_big_dictionary["Grade"].extend(grade_list)
                    the_big_dictionary["Grade"].append(" ")
                    the_big_dictionary["Exam Date"].extend([exam_date[table_index] for row in table.rows[2:]])
                    the_big_dictionary["Exam Date"].append(" ")
            else:   
                grades = [row.cells[2:] for row in table.rows]
                for i,sub_code in enumerate(sub_codes):
                    grade_list = []
                    for j,grade in enumerate(grades):
                        grade_list.append(grade[i].text)
                    the_big_dictionary["Subject Code"].extend([sub_code for row in table.rows])
                    the_big_dictionary["Subject Code"].append(" ")
                    the_big_dictionary["Register Number"].extend([row.cells[0].text for row in table.rows])
                    the_big_dictionary["Register Number"].append(" ")
                    the_big_dictionary["Student Name"].extend([row.cells[1].text.replace("\n"," ") for row in table.rows])
                    the_big_dictionary["Student Name"].append(" ")            
                    the_big_dictionary["Grade"].extend(grade_list)
                    the_big_dictionary["Grade"].append(" ")
                    the_big_dictionary["Exam Date"].extend([exam_date[table_index] for row in table.rows])
                    the_big_dictionary["Exam Date"].append(" ")
    return the_big_dictionary
#print_dictionary(convert_to_dictionary(r"result.docx","result release","Apr"))
pd.DataFrame(convert_to_dictionary(r"result.docx","Apr/May 2023")).to_excel("result.xlsx")
